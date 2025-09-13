from docx import Document
from docx.shared import Cm, Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import pytz
from config import WORD_FILE_MAIN, WORD_FILE_LIST

# Настройки форматирования
FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(12)
ALIGNMENT = 'center'
MOSCOW_TZ = pytz.timezone('Europe/Moscow')


class DocumentHandler:

    def __init__(self):
        self.main_doc = Document(WORD_FILE_MAIN)
        self.list_doc = Document(WORD_FILE_LIST)
    
    def get_keywords_fio(self):
        """Получение списка ФИО из основного документа"""
        return [self.main_doc.tables[0].cell(row, 2).text.lower() 
                for row in range(2, len(self.main_doc.tables[0].rows))]
    
    def create_fio_data_structure(self, keywords_fio):
        """Создание структуры данных для быстрого поиска по ФИО"""
        fio_data_list = []
        fio_by_last_name = {}
        
        for fio in keywords_fio:
            full_name = fio.lower()
            last_name = full_name.split(' ')[0]
            initials = full_name[:full_name.find(' ') + 2]
            fio_entry = {'full_name': full_name, 'last_name': last_name, 'initials': initials}
            fio_data_list.append(fio_entry)
            fio_by_last_name.setdefault(last_name, []).append(fio_entry)
        
        return fio_data_list, fio_by_last_name
    
    def get_today_date_info(self, vid=None, vremya='18:30-20:40'):
        """Получение информации о сегодняшней дате"""
        today = datetime.now(MOSCOW_TZ).strftime("%d-%m-%Y")
        chislo = today
        
        # Преобразование месяца
        month_dict = {
            '01': 'января', '02': 'февраля', '03': 'марта', '04': 'апреля',
            '05': 'мая', '06': 'июня', '07': 'июля', '08': 'августа',
            '09': 'сентября', '10': 'октября', '11': 'ноября', '12': 'декабря'
        }
        month = month_dict.get(chislo[3:5], '')
        
        # Обработка дня
        day = chislo[1] if chislo[0] == '0' else chislo[0:2]
        
        # Формирование строки времени
        yacheika = f'{vremya[:5]} {chislo[0:2]}.{chislo[3:5]}.{chislo[6:10]} г. - {vremya[6:]} {chislo[0:2]}.{chislo[3:5]}.{chislo[6:10]} г.'
        
        if vid is None:
            zagolovok = f"{' '.join(self.list_doc.paragraphs[2].text.split()[:-4])} {day} {month} {chislo[6:10]} г."
            return zagolovok, yacheika
        else:
            zagolovok = f'личного состава 3 курса 6 факультета, убывающего {vid} {day} {month} {chislo[6:10]} г.'
            return zagolovok, yacheika
    
    def sort_by_rank(self, table, ranks_fio):
        """Сортировка по воинским званиям"""
        rank_dict = {
            "старшина": 0,
            "старший сержант": 1,
            "сержант": 2,
            "младший сержант": 3,
            "рядовой": 4
        }
        
        sorted_rank_dict = sorted(ranks_fio, key=lambda item: rank_dict.get(item[0]))
        
        # Удаляем все строки, кроме первых двух
        for i in range(len(table.rows) - 2):
            table._element.remove(table.rows[-1]._element)
        
        for j, sort in enumerate(sorted_rank_dict):
            new_row = table.add_row()
            new_row.height = Cm(1.5) if len(sorted_rank_dict) % 9 == 7 else Cm(1.8)
            
            new_row.cells[0].text = f" {j + 1}. "
            for num_sort, cell in enumerate((0, 1, 2, 3, 5, 6)):
                yacheika = new_row.cells[cell]
                if cell != 0:
                    yacheika.text = sort[num_sort-1]
                self._format_cell(yacheika)
    
    def _format_cell(self, cell):
        """Форматирование ячейки"""
        for p in cell.paragraphs:
            for j in p.runs:
                j.font.name = FONT_NAME
                j.font.size = FONT_SIZE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt = p.paragraph_format
            fmt.space_after = Mm(0)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), ALIGNMENT)
            tcPr.append(vAlign)
    
    def insert_data_into_table(self, spisok_v_uval_fio, spisok_v_uval_forma, vid=None, vremya='18:30-20:40', mesto='г. Санкт-Петербург'):
        """Вставка данных в таблицу Word"""
        table = self.list_doc.tables[0]
        main_table = self.main_doc.tables[0]
        
        # Сортируем по ФИО
        spisok = sorted(spisok_v_uval_fio)
        ranks_fio = []
        
        # Поиск совпадений и заполнение данных
        for j, f in enumerate(spisok):
            for i in range(2, 90):  # 88 курсантов + 2 заголовочные строки
                main_table_list = main_table.cell(i, 2).text.upper()
                table_list = f.upper()
                if table_list == main_table_list:
                    zvanie = main_table.cell(i, 1).text
                    fio = main_table.cell(i, 2).text
                    telephone = main_table.cell(i, 3).text
                    forma = str(spisok_v_uval_forma[j][spisok_v_uval_forma[j].find('-')+1:spisok_v_uval_forma[j].find('Ф') - 1]).title()
                    
                    ranks_fio.append((zvanie, fio, telephone, forma, mesto))
        
        # Сортировка и вставка
        self.sort_by_rank(table, ranks_fio)
        length = len(ranks_fio) + 2
        
        # Обновление заголовка
        self.list_doc.paragraphs[2].text = self.get_today_date_info(vid, vremya)[0]
        self._format_paragraph(self.list_doc.paragraphs[2])
        
        # Обновление дат в ячейках
        for i in range(2, length):
            table.cell(i, 4).text = self.get_today_date_info(vid, vremya)[1]
            cell = table.cell(i, 4)
            self._format_cell(cell)
            first_paragraph = cell.paragraphs[0]
            first_paragraph.paragraph_format.first_line_indent = Cm(0.24)
        
        self.list_doc.save(WORD_FILE_LIST)
    
    def _format_paragraph(self, paragraph):
        """Форматирование абзаца"""
        for i in paragraph.runs:
            i.font.name = FONT_NAME
            i.font.size = FONT_SIZE
    
    def change_officer(self, doljnost, zvanie, oficer_fio):
        """Замена офицера в документе"""
        # Редактирование должности и звания офицера
        self.list_doc.paragraphs[5].text = doljnost
        self.list_doc.paragraphs[6].text = zvanie
        self.list_doc.paragraphs[7].text = oficer_fio
        
        for paragraph in [5, 6, 7]:
            self._format_paragraph(self.list_doc.paragraphs[paragraph])
        
        self.list_doc.save(WORD_FILE_LIST)
    
    def clean_table(self):
        """Очистка таблицы"""
        table = self.list_doc.tables[0]
        
        for i in range(len(table.rows) - 2):
            table._element.remove(table.rows[-1]._element)
        
        self.list_doc.save(WORD_FILE_LIST)
    
    def update_date_header(self):
        """Обновление заголовка с датой"""
        self.list_doc.paragraphs[2].text = self.get_today_date_info()[0]
        self.list_doc.save(WORD_FILE_LIST)
    
    def get_responsible_officer(self):
        """Получение информации об ответственном офицере"""
        return f"{self.list_doc.paragraphs[6].text} {self.list_doc.paragraphs[7].text}"
    
    def get_document_date(self):
        """Получение даты из документа"""
        return self.list_doc.paragraphs[2].text.split()[-4:]